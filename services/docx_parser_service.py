from docx import Document
import docx2txt

def extract_clean_text_from_docx(file_path: str) -> str:
    def extract_with_python_docx(path: str) -> str:
        try:
            doc = Document(path)
            lines = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        clean = cell.text.strip().replace("\n", " ")
                        if clean:
                            row_text.append(clean)
                    lines.extend(row_text)
            return "\n".join(lines).strip()
        except Exception:
            return ""

    def extract_with_docx2txt(path: str) -> str:
        try:
            return docx2txt.process(path).strip()
        except:
            return ""

    text = extract_with_python_docx(file_path)
    if len(text.split()) < 30:
        return extract_with_docx2txt(file_path)
    return text
